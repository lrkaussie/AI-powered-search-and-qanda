"""Service for processing and extracting content from various document formats."""

from pathlib import Path
from typing import Any

import fitz  # type: ignore  # PyMuPDF
from docx import Document as DocxDocument

from app.core.exceptions import DocumentProcessingError
from app.models.document import Document


class DocumentProcessor:
    """Service for processing different types of documents."""

    supported_formats = {".pdf", ".docx", ".txt"}

    @staticmethod
    async def process_document(file_path: Path) -> Document:
        """Process a document file and extract its content.

        Args:
            file_path: Path to the document file

        Returns:
            Document: Document object with extracted content and metadata

        Raises:
            DocumentProcessingError: If there's an error processing the document
        """
        try:
            file_extension = file_path.suffix.lower()

            if file_extension not in DocumentProcessor.supported_formats:
                supported_formats_str = ", ".join(DocumentProcessor.supported_formats)
                msg = "Unsupported file format. "
                msg += f"Supported formats: {supported_formats_str}"
                raise DocumentProcessingError(msg)

            content = ""
            metadata: dict[str, Any] = {}

            # Process based on file type
            if file_extension == ".pdf":
                content, metadata = await DocumentProcessor._process_pdf(file_path)
            elif file_extension == ".docx":
                content, metadata = await DocumentProcessor._process_docx(file_path)
            elif file_extension == ".txt":
                content = await DocumentProcessor._process_txt(file_path)

            # Create document object
            document = Document(
                id=str(hash(file_path)),
                title=str(file_path.stem),
                content=content,
                doc_type=file_extension[1:],  # Remove the dot
                metadata=metadata,
            )

            return document

        except Exception as e:
            raise DocumentProcessingError(f"Error processing document: {str(e)}")

    @staticmethod
    async def _process_pdf(file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process PDF files using PyMuPDF.

        Args:
            file_path: Path to the PDF file

        Returns:
            tuple[str, dict[str, Any]]: Tuple containing:
                - str: Extracted text content
                - dict: Document metadata

        Raises:
            DocumentProcessingError: If there's an error processing the PDF
        """
        try:
            text_content = []
            metadata = {}

            # Extract text with PyMuPDF
            with fitz.open(file_path) as pdf:
                metadata = {
                    "page_count": len(pdf),
                    "title": pdf.metadata.get("title", ""),
                    "author": pdf.metadata.get("author", ""),
                    "subject": pdf.metadata.get("subject", ""),
                }

                for page in pdf:
                    text_content.append(page.get_text())

            return "\n".join(text_content), metadata

        except Exception as e:
            raise DocumentProcessingError(f"Error processing PDF: {str(e)}")

    @staticmethod
    async def _process_docx(file_path: Path) -> tuple[str, dict[str, Any]]:
        """Process DOCX files.

        Args:
            file_path: Path to the DOCX file

        Returns:
            tuple[str, dict[str, Any]]: Tuple containing:
                - str: Extracted text content
                - dict: Document metadata

        Raises:
            DocumentProcessingError: If there's an error processing the DOCX
        """
        try:
            doc = DocxDocument(str(file_path))
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            content = "\n".join(paragraphs)

            metadata = {
                "core_properties": {
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created,
                    "modified": doc.core_properties.modified,
                }
            }

            return content, metadata

        except Exception as e:
            raise DocumentProcessingError(f"Error processing DOCX: {str(e)}")

    @staticmethod
    async def _process_txt(file_path: Path) -> str:
        """Process plain text files.

        Args:
            file_path: Path to the text file

        Returns:
            str: Content of the text file

        Raises:
            DocumentProcessingError: If there's an error processing the text file
        """
        try:
            with open(file_path, encoding="utf-8") as f:
                return f.read()

        except Exception as e:
            raise DocumentProcessingError(f"Error processing TXT: {str(e)}")
